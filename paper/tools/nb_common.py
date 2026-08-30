"""Shared helpers for the working-notebook generator (python-docx)."""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_BREAK


class NB:
    def __init__(self):
        self.d = Document()
        self.d.styles["Normal"].font.name = "Calibri"
        self.d.styles["Normal"].font.size = Pt(10.5)

    def h(self, t, lvl=1):
        self.d.add_heading(t, lvl)

    def title(self, t):
        self.d.add_heading(t, 0)

    def p(self, t="", bold=False, italic=False):
        para = self.d.add_paragraph(); r = para.add_run(t); r.bold = bold; r.italic = italic; return para

    def bl(self, items, style="List Bullet"):
        for it in items:
            self.d.add_paragraph(it, style=style)

    def num(self, items):
        self.bl(items, style="List Number")

    def lines(self, n=4):
        for _ in range(n):
            self.d.add_paragraph("_" * 95)

    def tbl(self, rows):
        t = self.d.add_table(rows=len(rows), cols=len(rows[0])); t.style = "Table Grid"
        for i, row in enumerate(rows):
            for j, cell in enumerate(row):
                c = t.cell(i, j); c.text = str(cell)
                if i == 0:
                    for r in c.paragraphs[0].runs:
                        r.bold = True
        self.d.add_paragraph()

    def pagebreak(self):
        self.d.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    def box(self, title, text):
        self.p(title, bold=True); self.p(text, italic=True)

    def experiment(self, title, *, background, question, design, says, not_shown, where, status, gaps=None):
        """The fixed per-experiment template."""
        self.h(title, 2)
        self.p("Background and motivation", bold=True); self.p(background)
        self.p("Question it answers", bold=True); self.p(question)
        self.p("Design", bold=True); self.bl(design) if isinstance(design, list) else self.p(design)
        self.p("What the figure / table says", bold=True); self.bl(says) if isinstance(says, list) else self.p(says)
        self.p("What it does not show / caveats", bold=True); self.bl(not_shown) if isinstance(not_shown, list) else self.p(not_shown)
        self.p("Where it lives (script → data → figure)", bold=True); self.p(where, italic=True)
        self.p("Status", bold=True); self.p(status)
        if gaps:
            self.p("Gaps for you to fill", bold=True); self.bl(gaps)
        self.p("Your notes:", bold=True); self.lines(3)

    def save(self, path):
        self.d.save(path)
