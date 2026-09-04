"""Rebuild P1_core_experiments.docx as a design-only document (PI order
2026-09-01): no prose, only experiment design; Exp 4 replaced by the
CPU-Drake-CENIC vs GPU-Newton-CENIC per-world scaling comparison; no
matplotlib scene renders (simulator-viewer captures pending)."""

import os

from docx import Document
from docx.shared import Inches

FIG = os.path.expanduser("~/Documents/code/icra2027/part1/bench/results/figures")
OUT = os.path.expanduser("~/Documents/code/icra2027/part1/bench/results/P1_core_experiments.docx")

H1 = "91c8394"
H0 = "91c8394"
H3 = "91c8394"
B3 = f"https://github.com/mardigiorgio/icra2027/blob/{H3}"
B1 = f"https://github.com/mardigiorgio/icra2027/blob/{H1}"
B0 = f"https://github.com/mardigiorgio/icra2027/blob/{H0}"

d = Document()

# Freshness gate: every figure must postdate the 8-body scene switch
# (2026-09-03 20:45 local); an older file means its rerun is still in flight
# and the doc says so instead of showing a stale 20-body figure.
import datetime as _dt
_FRESH = _dt.datetime(2026, 9, 3, 20, 45).timestamp()


def _pic(name, width):
    path = os.path.join(FIG, name)
    if os.path.exists(path) and os.path.getmtime(path) >= _FRESH:
        d.add_picture(path, width=width)
    else:
        d.add_paragraph(f"[figure {name} pending: its rerun on the 8-body scenes is in progress and replaces this line when done]")

d.add_heading("Part 1, core experiments (design)", level=1)
d.add_paragraph(f"Build {_dt.datetime.now():%Y-%m-%d %H:%M}. Every figure and number below is from the 8-body clutter scenes (Marco's ruling of 2026-09-03), measured on 2026-09-03 with icf_warp perf/contact-solve 0fca334 (Newton/Warp) and Drake 1.56.0 (CPU reference). Code citations point at icra2027 commit 91c8394, the code these runs used.")
p = d.add_paragraph()
r = p.add_run("Design-only revision, 2026-09-01, on the PI's request: each experiment lists scene, arms, what is varied, what is measured, and code. Scene images are captures from Newton's own viewer (ViewerGL): what is shown is what the simulator draws.")
r.italic = True

d.add_paragraph("Reading key:", style=None)
for t in [
    "Two contact solvers: MuJoCo (solref constraint contact) and ICF (convex contact with a literal stiffness k in N/m)",
    "Two stepping modes each: FIXED (one chosen step size δt) and ERROR CONTROL, EC (step-doubling estimate, per-world step adapts so position error stays below a requested tolerance ε)",
    "CENIC = ICF + error control",
    "Smaller δt or smaller ε: more accurate, more expensive",
    "Contact fairness: every arm, MuJoCo included, consumes the SAME Newton collision pipeline contact set; the solver is the only difference between arms",
    "Clutter worlds: every parallel world draws its own initial arrangement (per-world seed, rejection-sampled so no bodies overlap at spawn; the same shared generator feeds the Drake harness, so CPU and GPU integrate identical world sets)",
]:
    d.add_paragraph(t, style="List Bullet")

# ------------------------------------------------------------------ Exp 1
d.add_heading("Experiment 1, Contact stiffness realization: penetration vs step and vs tolerance", level=2)
_pic("capture_stiffness.png", Inches(3.2))
d.add_paragraph("Scene (Newton viewer): the 65 g sphere at rest on the plane.")
_pic("stiffness_sweep.png", Inches(6.2))
d.add_paragraph("Figure: resting penetration of the sphere divided by the model's m g / k (1 = the stiffness the scene asked for). (a) fixed step, both solvers; the dotted vertical line is where MuJoCo's refsafe clamp starts to bite (δt = τ/2 = 1.2 ms), beyond which its contact time constant is 2 δt instead of τ. (b) error control at the requested tolerance ε.")
d.add_paragraph("Design:")
for t in [
    "Scene: one sphere, r = 2.5 cm, density 1000 (65 g), resting on a flat floor; friction 0.5; contact margin 0; 3 s settle before readout",
    "Contact stiffness: k = 100,000 N/m for the whole experiment (companion rows at k = 1,000 in the same CSV); ICF takes k directly. MuJoCo takes the reference solref (timeconst τ = 2.4 ms, dampratio 1), τ calibrated so the converged rest depth equals m g / k at a 1 ms step; MuJoCo clamps timeconst to at least 2 δt (refsafe), so above δt = τ/2 the contact is softer than asked by (2 δt / τ)^2. The direct solref form (-k, -b), a literal stiffness the clamp does not touch, is run as a control: it is unstable at 10 and 5 ms and exact at 2 ms and below (rows in the CSV, not on the figure)",
    "Arms: ICF fixed, MuJoCo fixed, ICF EC (CENIC), MuJoCo EC",
    "Varied, panel (a): fixed step δt in {10, 5, 2, 1, 0.5} ms; varied, panel (b): tolerance ε in {0.1 ... 0.00001} m; nothing else changes between runs",
    "Measured: resting penetration divided by mg/k (1 = the model); a launched sphere is marked unstable",
    "Built-in control: at δt <= 1 ms both solvers read 1.00 within 2 percent (MuJoCo 0.98 at 1 ms and 0.5 ms), which is the definition of same physical system (equality by convergence, no parameter translation anywhere)",
    "Result (2026-09-03 run): ICF reads 1.00 at every step and every tolerance. MuJoCo fixed reads 57x at 10 ms, 16.8x at 5 ms, 2.7x at 2 ms, 0.98 at 1 ms and 0.5 ms. MuJoCo error control reads 16.8x for every ε from 0.1 to 1e-4 and 1.25x at 1e-5: its controller only brings the step under τ/2 at the tightest tolerance, because the softened contact is self-consistent at every step and produces no error signal. At k = 1,000 (τ = 31.8 ms) both solvers read 1.00 at every step; the clamp never bites",
]:
    d.add_paragraph(t, style="List Bullet")
d.add_paragraph("Code:")
for t in [
    f"step and tolerance grids, {B1}/part1/bench/benchmarks/part1_stiffness_sweep.py#L40-L41",
    f"the calibrated MuJoCo solref anchor, self-verifying at fine steps, {B1}/part1/bench/benchmarks/part1_stiffness_sweep.py#L43-L47",
    f"MuJoCo solref calibration anchors, {B1}/part1/scenes/cenic_scenes.py#L66-L69",
    f"scene build, settle loop and depth readout, {B1}/part1/bench/benchmarks/part1_stiffness_sweep.py#L50-L79",
    f"the sweep cells (arms x steps x tolerances), {B1}/part1/bench/benchmarks/part1_stiffness_sweep.py#L86-L103",
    f"solref written onto every contact geom, {B1}/part1/bench/four_arms.py#L103-L113",
]:
    d.add_paragraph(t, style="List Bullet")

# ------------------------------------------------------------------ Exp 2
d.add_heading("Experiment 2, Work-precision on the clutter scenes", level=2)
_pic("capture_soft_clutter.png", Inches(5.4))
d.add_paragraph("Soft clutter (8 spheres), world 0 (Newton viewer): t = 0 left, settled at 1.5 s right.")
_pic("capture_hard_clutter.png", Inches(5.4))
d.add_paragraph("Hard clutter (4 spheres + 4 cubes), world 0 (Newton viewer): t = 0 left, settled at 1.5 s right.")
_pic("workprecision.png", Inches(6.2))
d.add_paragraph("Figure: wall time to complete the 2 s clutter scene, vs requested tolerance; dotted levels are the fixed-step arms at δt = 10 ms and 1 ms; the thin gray line is real time (wall equal to simulated time). Any timed-out or budget-limited cell would be drawn as a cross; none occurred.")
d.add_paragraph("Design:")
for t in [
    "Scenes (Marco's ruling 2026-09-03: two lattice layers, 8 bodies, since per-world contact cost sets every axis here): soft clutter (8 spheres, k = 1,000 N/m) and hard clutter (4 spheres + 4 cubes, k = 100,000 N/m), dropped into a 30 cm bin; each world its own randomized non-overlapping lattice (the first 8 bodies of the original 20-body stream); 2 s horizon",
    "Arms: ICF EC and MuJoCo EC (curves); ICF fixed and MuJoCo fixed at 10 ms and 1 ms (dotted reference levels)",
    "Varied: requested tolerance ε in {0.1 ... 0.000001} m; number of parallel worlds N in {1, 1024}; every cell is the median of 3 independent runs",
    "Measured: wall time to complete the 2 s scene; timeout at 100 s per world-second and march-budget exhaustion are marked as crosses",
    "Held: scene geometry, masses, friction, δt_max = 0.1 s, march budget of 16384 substeps per 0.1 s boundary for every arm (raised from 4096 on 2026-09-03 after the 20-body hard clutter at ε = 1e-6 needed between 4096 and 16384 substeps in its impact phase; on the 8-body scenes no cell of any arm exhausts the budget at any tolerance)",
    "Status (2026-09-03, evening): both rows measured on the 8-body scenes with the ICF arms on icf_warp perf/contact-solve 0fca334 at the 16384 march budget; the 20-body results of the same day are kept as *_20.csv",
]:
    d.add_paragraph(t, style="List Bullet")
d.add_paragraph("Code:")
for t in [
    f"timeout semantics (per world-second), {B0}/part1/bench/benchmarks/part1_workprecision.py#L6-L16",
    f"accuracy ladder + march budget, {B0}/part1/bench/benchmarks/part1_workprecision.py#L18-L24",
    f"first attempt k_Init·δt_max, {B0}/part1/bench/four_arms.py#L47",
    f"inner tolerance rule (CENIC eq. 34), {B0}/part1/bench/four_arms.py#L49-L50",
    f"contact budgets (>= 2x measured demand), {B0}/part1/bench/four_arms.py#L55-L65",
    f"per-world no-overlap lattice generator, {B3}/part1/scenes/clutter_lattice.py#L38-L86",
    f"MuJoCo arms on Newton pipeline contacts, {B3}/part1/bench/four_arms.py#L124-L131 and #L157-L162",
]:
    d.add_paragraph(t, style="List Bullet")

# ------------------------------------------------------------------ Exp 3
d.add_heading("Experiment 3, Wall time vs number of parallel worlds", level=2)
_pic("capture_hard_clutter.png", Inches(5.4))
d.add_paragraph("Hard clutter (4 spheres + 4 cubes), world 0 (Newton viewer): t = 0 left, settled at 1.5 s right; every world of the batch is its own randomized lattice of this scene.")
_pic("scaling_hard-clutter.png", Inches(6.2))
d.add_paragraph("Figure: wall time to complete the 2 s hard clutter scene, vs the number of parallel worlds. Each point is 20 times the median per-boundary wall of one run over its 20 timed 0.1 s boundaries; the band spans the median to the 90th percentile of those per-boundary walls.")
d.add_paragraph("Design:")
for t in [
    "Scene: hard clutter (as in Experiment 2)",
    "Arms: all four; fixed arms at δt = 10 ms, EC arms at ε = 0.001",
    "Varied: number of parallel worlds N = 2^6 ... 2^13; nothing else",
    "Measured: per-boundary wall over the 20 timed boundaries of one run per cell; the plotted scene time is 20 x the median boundary, the band the median-to-p90 spread of the boundaries",
    "Held: scene, tolerance, timed window (2 s after 0.2 s warm-up)",
    "Status (2026-09-03, evening): all four arms measured on the 8-body hard clutter with the ICF arms on icf_warp perf/contact-solve 0fca334; the 20-body results of the same day are kept as part1_scaling_*_20.csv",
]:
    d.add_paragraph(t, style="List Bullet")
d.add_paragraph("Code:")
for t in [
    f"runner (per-N subprocess), {B0}/part1/bench/benchmarks/part1_scaling.py#L36",
    f"warm-up handling, {B0}/part1/bench/benchmarks/part1_scaling.py#L44-L47",
    f"median statistic, {B0}/part1/bench/benchmarks/part1_scaling.py#L72",
    f"defaults: 2 s timed, 0.2 s warm-up, δt_fixed = 10 ms, ε = 1e-3, {B0}/part1/bench/benchmarks/part1_scaling.py#L80-L86",
]:
    d.add_paragraph(t, style="List Bullet")

# ------------------------------------------------------------------ Exp 4 (new)
H2 = "91c8394"
B2 = f"https://github.com/mardigiorgio/icra2027/blob/{H2}"
d.add_heading("Experiment 4, CENIC reference implementation (Drake, CPU) vs this work (Newton/Warp, GPU): wall time vs number of worlds", level=2)
_pic("capture_hard_clutter.png", Inches(5.4))
d.add_paragraph("Hard clutter (4 spheres + 4 cubes), world 0 (Newton viewer): t = 0 left, settled at 1.5 s right; the CPU reference integrates the identical randomized worlds (same lattice generator, same seeds).")
_pic("cenic_scaling.png", Inches(6.2))
d.add_paragraph("Figure: wall time for the batch of N worlds to complete the 2 s scene (warm-up to 0.2 s excluded, one makespan over the 20 timed control boundaries, identical protocol on both stacks), CPU reference (128 cores, 96 workers, worlds beyond the cores in sequence) vs this work (RTX 5090), both at requested accuracy 0.001. Scene: the hard clutter of Experiment 2 (8 bodies), both stacks drawing the same bodies from the same seeded stream.")
d.add_paragraph("Design (replaces the actuated PD push, PI 2026-09-01):")
for t in [
    "Question: how long a batch of N worlds takes to complete the scene under error-controlled CENIC, reference CPU implementation vs the GPU implementation this paper contributes",
    "CPU arm: Drake 1.56.0, integration_scheme = cenic (the reference implementation, shipped in Drake); W = min(N, 96) worker processes on the idle 128-core host, each building ONE diagram and hosting its share of worlds as ~1 MB per-world contexts, so worlds beyond the core count run SEQUENTIALLY on the cores available with no artificial per-world residency cost",
    "Lockstep batch semantics in both stacks: every worker advances each of its worlds one 0.1 s boundary, then all workers barrier; all workers start on one shared GO after building and warming up; the parent times one makespan",
    "GPU arm: Newton/Warp CENIC (ICF EC), batched worlds, one fair-protocol cell per N so both stacks share every x point",
    "Scene: the hard clutter of Experiment 2 (8 bodies: 4 spheres + 4 cubes, part1/scenes/clutter_lattice.py `layers` = 2). Both stacks import the same lattice generator, world i seeded identically, so CPU and GPU integrate identical randomized world sets",
    "Varied: number of concurrent worlds, 1 ... 16384 (fourteen doublings); the CPU is past its 128 cores from N = 128; the GPU is still sublinear at 16384",
    "Measured: plain wall-clock time for the batch to complete the 2 s scene, in seconds; no normalization; lower is better",
    "Held in both stacks: max step 0.1 s, 0.2 s warm-up then 2 s timed, contact constants (k = 100,000 N/m, Hunt-Crossley 1.0 s/m, friction 0.5, stiction tolerance 0.0001); requested tolerance 0.001 in both stacks for the nominal curves",
    "Tolerance note: both stacks run at their requested accuracy 0.001. Drake's `accuracy` is a relative tolerance on its weighted state norm; this work's eps is an absolute position L-inf tolerance in metres, so the two controllers take different step counts on the same world (Drake 543 vs this work 105 accepted steps over the 20 timed boundaries of world 0, both reaching the same rest state); the figure compares the two implementations at their own nominal 0.001",
    "Small-N protocol and trials (2026-09-03): N <= 8 on both stacks are medians of repeated runs (GPU 5 trials: N = 1 0.084 to 0.111 s, N = 2 0.126 to 0.140, N = 4 0.147 to 0.199, N = 8 0.187 to 0.235; CPU 3 trials: N = 1 0.085 to 0.089, N = 2 0.101 to 0.106, N = 4 0.101 to 0.106, N = 8 0.107 to 0.115, N = 16 0.161 to 0.163, N = 32 0.185 to 0.188, N = 64 0.201 to 0.213). When every world has its own worker (N <= 96) the CPU makespan is the slowest world's integration time, because the lockstep barrier there only synchronizes clocks and its file polling (about 55 ms per run) is a harness artifact, not the reference implementation; with more worlds than workers the barrier makespan is the batch time. The barrier makespan is kept alongside in the CSV",
    "Measured result (seconds to complete the scene, CPU / GPU): 1: 0.0852 / 0.096; 2: 0.101 / 0.128; 4: 0.102 / 0.155; 8: 0.108 / 0.196; 16: 0.161 / 0.3; 32: 0.188 / 0.3; 64: 0.202 / 0.4; 128: 0.472 / 0.5; 256: 0.691 / 0.7; 512: 1.38 / 0.8; 1024: 2.4 / 1; 2048: 4.44 / 1.3; 4096: 8.58 / 1.9; 8192: 16.5 / 3; 16384: 33 / 5.5. Below about 256 worlds the CPU is faster (at one world the GPU's 0.096 s against Drake's 0.085 s; per step the GPU is about six times slower there); the curves cross between 256 and 512 worlds; the GPU is 6.0x faster at 16384 and still sublinear there (exponent 4096 to 16384: GPU 0.77 vs CPU 0.97); saturated per-world cost GPU 0.34 ms vs CPU 2.0 ms",
    "Solver state for these numbers: icf_warp_adaptive branch perf/contact-solve at 0fca334 (block-structured assembly on the free-body layout, 24-dof block-sparse Cholesky, narrow-tail march compaction ICF_MARCH_COMPACT=1, exact per-iteration launch folds); physics unchanged (64-world oracle inside the run-to-run envelope, zero solve failures). Before this work the GPU arm cost about 30 ms per added world on the 20-body scene, 2.1x the CPU, with no crossover",
    "Reading: a single world's march is a serial chain on either machine, and one CPU core runs it faster than the GPU (per step about 6x); the GPU's advantage appears only once enough worlds run concurrently to fill the card, which on this scene happens between 256 and 512 worlds, and it grows from there because the GPU adds worlds at a fraction of a millisecond each while the CPU runs them in sequence on its cores. The GPU case additionally rests on co-residence with the learner (no PCIe state transfer in the training loop)",
]:
    d.add_paragraph(t, style="List Bullet")
d.add_paragraph("Code:")
for t in [
    f"shared lattice import and scene constants, {B3}/part1/bench/benchmarks/part1_cenic_cpu.py#L41-L59",
    f"Drake world build and CENIC configuration, {B3}/part1/bench/benchmarks/part1_cenic_cpu.py#L70-L150",
    f"GO-signal barrier, lockstep boundaries, makespan timing, {B3}/part1/bench/benchmarks/part1_cenic_cpu.py#L157-L205",
    f"GPU arm on the same protocol (warm-up, timed window, makespan), {B0}/part1/bench/benchmarks/part1_gpu_fair.py#L26-L28 and #L41-L90",
    f"Both ladders, one row per N with the solver commit, {B0}/part1/bench/results/part1_gpu_fair_ladder.csv and {B0}/part1/bench/results/part1_cenic_cpu.csv",
]:
    d.add_paragraph(t, style="List Bullet")

d.save(OUT)
print("wrote", OUT)
doc = Document(OUT)
full = "\n".join(p.text for p in doc.paragraphs)
assert full.count(chr(8212)) == 0, "em dash found"
print("paragraphs:", len(doc.paragraphs), "images:", len(doc.inline_shapes))
